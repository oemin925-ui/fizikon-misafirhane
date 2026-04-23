from pathlib import Path

from docx import Document


base = Path(r"C:\Users\ASUS\Documents\Codex\2026-04-22-files-mentioned-by-the-user-2")

for path in sorted(base.glob("*.docx")):
    print(f"FILE: {path.name}")
    doc = Document(str(path))
    print("PARAGRAPHS:")
    shown = 0
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue
        shown += 1
        print(f"  P{i}: {text}")
        if shown >= 10:
            break
    print(f"TABLES: {len(doc.tables)}")
    for table_index, table in enumerate(doc.tables, 1):
        print(f"  TABLE {table_index}: rows={len(table.rows)} cols~={len(table.columns)}")
        for row_index, row in enumerate(table.rows[:10], 1):
            values = [cell.text.replace('\n', ' / ').strip() for cell in row.cells[:10]]
            print(f"    R{row_index}: {values}")
    print("-" * 60)
